"""
OMNI-LEARN OS — 大型PDF电子书阅读与知识提取工具
pdf_reader_tool.py

功能：
  - 本地读取任意大小的PDF文件（无需上传）
  - 自动提取目录结构
  - 逐章总结知识点
  - 提取核心词汇
  - 生成练习题
  - 导出Word/Excel笔记

使用：
  streamlit run pdf_reader_tool.py

依赖：
  pip install pymupdf streamlit openai python-dotenv openpyxl python-docx
"""

import streamlit as st
import os
import sys
import json
import time
import math
import tempfile
from pathlib import Path

from dotenv import load_dotenv
load_dotenv()

from openai import OpenAI

st.set_page_config(
    page_title="OMNI PDF 电子书工具",
    page_icon="📚",
    layout="wide"
)

st.markdown("""
<style>
.chapter-card { background:white; border-radius:10px; padding:14px 18px;
  margin-bottom:10px; border-left:4px solid #3b82f6;
  box-shadow:0 2px 8px rgba(0,0,0,0.06); }
.summary-box { background:#f0f9ff; border-radius:10px; padding:16px 20px;
  margin-bottom:12px; border:1px solid #bae6fd; line-height:1.8; }
.vocab-pill { display:inline-block; background:#fef3c7; color:#92400e;
  padding:3px 10px; border-radius:12px; font-size:.85rem; margin:3px;
  font-weight:500; }
.progress-info { background:#f0fdf4; border-radius:8px; padding:10px 14px;
  font-size:.88rem; color:#166534; margin-bottom:10px; }
</style>
""", unsafe_allow_html=True)

# ── 客户端 ──────────────────────────────────────
@st.cache_resource
def get_client():
    api_key = os.getenv("DEEPSEEK_API_KEY", "")
    if not api_key:
        st.error("❌ 请在 .env 文件写入 DEEPSEEK_API_KEY")
        st.stop()
    return OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

client = get_client()

# ── Session State ────────────────────────────────
for k, v in {
    "pdf_path": "",
    "pdf_name": "",
    "total_pages": 0,
    "chapters": [],
    "summaries": {},
    "vocab_list": {},
    "questions": {},
    "processing": False,
    "current_chapter": None,
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══════════════════════════════════════════════════
# PDF 处理函数
# ══════════════════════════════════════════════════
def load_pdf(path: str):
    """加载 PDF，返回页面文本列表"""
    try:
        import fitz  # pymupdf
        doc = fitz.open(path)
        pages = []
        for page in doc:
            text = page.get_text()
            pages.append(text)
        doc.close()
        return pages
    except ImportError:
        st.error("请安装 pymupdf：pip install pymupdf")
        return []
    except Exception as e:
        st.error(f"PDF 读取失败：{e}")
        return []

def extract_toc(path: str) -> list:
    """提取 PDF 目录（如果有）"""
    try:
        import fitz
        doc = fitz.open(path)
        toc = doc.get_toc()
        doc.close()
        return toc  # [(level, title, page), ...]
    except:
        return []

def pages_to_chunks(pages: list, chunk_size: int = 8) -> list:
    """将页面列表分成块"""
    chunks = []
    for i in range(0, len(pages), chunk_size):
        chunk_pages = pages[i:i+chunk_size]
        chunk_text = "\n\n".join(chunk_pages)
        chunks.append({
            "start_page": i + 1,
            "end_page": min(i + chunk_size, len(pages)),
            "text": chunk_text[:4000],  # 限制每块长度
        })
    return chunks

def summarize_chunk(text: str, book_name: str, lang: str = "zh") -> dict:
    """用 DeepSeek 总结一块内容"""
    if lang == "zh":
        prompt = f"""请阅读以下来自《{book_name}》的内容，用中文给出：
1. 核心知识点（3-5条）
2. 重要词汇（5-8个英文词汇，含中文释义）
3. 实用例句（2-3句）

内容：
{text[:3000]}

以JSON格式返回：
{{
  "key_points": ["知识点1", "知识点2", "知识点3"],
  "vocabulary": [{{"word": "word1", "cn": "释义", "example": "例句"}}],
  "useful_phrases": ["短语1", "短语2"]
}}"""
    else:
        prompt = f"""Read the following content from "{book_name}" and provide:
1. Key learning points (3-5 bullets)
2. Important vocabulary (5-8 words with Chinese meanings)
3. Useful phrases (2-3)

Content: {text[:3000]}

Return JSON:
{{"key_points": [...], "vocabulary": [...], "useful_phrases": [...]}}"""

    try:
        r = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.3
        )
        raw = r.choices[0].message.content.strip()
        raw = raw.replace("```json", "").replace("```", "").strip()
        return json.loads(raw)
    except Exception as e:
        return {
            "key_points": [f"内容解析中... ({str(e)[:50]})"],
            "vocabulary": [],
            "useful_phrases": []
        }

def generate_questions_for_chunk(text: str, book_name: str) -> list:
    """为内容块生成练习题"""
    prompt = f"""Based on this content from "{book_name}", create 3 practice questions.

Content: {text[:2000]}

Return ONLY JSON array:
[
  {{"type": "multiple_choice", "question": "Q?",
    "options": ["A. ...", "B. ...", "C. ...", "D. ..."],
    "answer": "A", "explanation": "Because..."}},
  {{"type": "fill_blank", "question": "The ___ is important in business.",
    "answer": "key word", "explanation": "..."}},
  {{"type": "true_false", "question": "Statement here.",
    "answer": "True", "explanation": "..."}}
]"""

    try:
        r = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=600,
            temperature=0.4
        )
        raw = r.choices[0].message.content.strip()
        raw = raw.replace("```json", "").replace("```", "").strip()
        return json.loads(raw)
    except:
        return []

# ══════════════════════════════════════════════════
# 导出函数
# ══════════════════════════════════════════════════
def export_to_word(book_name: str, summaries: dict, vocab_list: dict) -> str:
    """导出 Word 笔记"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(f"📚 {book_name} — 学习笔记", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"由 OMNI-LEARN OS AI 自动生成 | {time.strftime('%Y-%m-%d')}")
        doc.add_paragraph("")

        for chunk_key, summary in summaries.items():
            doc.add_heading(f"📖 {chunk_key}", level=1)

            if summary.get("key_points"):
                doc.add_heading("核心知识点", level=2)
                for pt in summary["key_points"]:
                    p = doc.add_paragraph(f"• {pt}", style="List Bullet")

            if summary.get("vocabulary"):
                doc.add_heading("重要词汇", level=2)
                for v in summary["vocabulary"]:
                    w = v.get("word","")
                    cn = v.get("cn","")
                    ex = v.get("example","")
                    doc.add_paragraph(f"🔑 {w} — {cn}")
                    if ex:
                        p = doc.add_paragraph(f"   例句：{ex}")
                        p.runs[0].italic = True

            if summary.get("useful_phrases"):
                doc.add_heading("实用短语", level=2)
                for ph in summary["useful_phrases"]:
                    doc.add_paragraph(f"💬 {ph}")

            doc.add_paragraph("")

        # 词汇汇总
        all_vocab = []
        for s in summaries.values():
            all_vocab.extend(s.get("vocabulary", []))

        if all_vocab:
            doc.add_heading("📋 词汇总表", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "英文单词"
            hdr[1].text = "中文释义"
            hdr[2].text = "例句"
            for v in all_vocab:
                row = table.add_row().cells
                row[0].text = v.get("word", "")
                row[1].text = v.get("cn", "")
                row[2].text = v.get("example", "")

        out_path = os.path.join(tempfile.gettempdir(),
                                f"{book_name[:20]}_笔记.docx")
        doc.save(out_path)
        return out_path
    except ImportError:
        st.error("请安装 python-docx：pip install python-docx")
        return ""

def export_to_excel(book_name: str, summaries: dict) -> str:
    """导出 Excel 词汇表"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = "词汇表"

        # Headers
        headers = ["章节/页码", "英文单词", "中文释义", "例句", "实用短语"]
        for ci, h in enumerate(headers, 1):
            c = ws.cell(row=1, column=ci, value=h)
            c.font = Font(bold=True, color="FFFFFF", name="Arial")
            c.fill = PatternFill("solid", fgColor="1F3864")
            c.alignment = Alignment(horizontal="center", vertical="center")

        row = 2
        for chunk_key, summary in summaries.items():
            for v in summary.get("vocabulary", []):
                ws.cell(row=row, column=1, value=chunk_key)
                ws.cell(row=row, column=2, value=v.get("word",""))
                ws.cell(row=row, column=3, value=v.get("cn",""))
                ws.cell(row=row, column=4, value=v.get("example",""))
                phrases = "; ".join(summary.get("useful_phrases", []))
                ws.cell(row=row, column=5, value=phrases if row == 2 else "")
                row += 1

        # Sheet 2: Key Points
        ws2 = wb.create_sheet("知识点总结")
        ws2.cell(row=1, column=1, value="章节").font = Font(bold=True)
        ws2.cell(row=1, column=2, value="核心知识点").font = Font(bold=True)
        row2 = 2
        for chunk_key, summary in summaries.items():
            for pt in summary.get("key_points", []):
                ws2.cell(row=row2, column=1, value=chunk_key)
                ws2.cell(row=row2, column=2, value=pt)
                row2 += 1

        for col in ["A","B","C","D","E"]:
            ws.column_dimensions[col].width = 20
        ws2.column_dimensions["A"].width = 15
        ws2.column_dimensions["B"].width = 60

        out_path = os.path.join(tempfile.gettempdir(),
                                f"{book_name[:20]}_词汇表.xlsx")
        wb.save(out_path)
        return out_path
    except Exception as e:
        st.error(f"Excel 导出失败：{e}")
        return ""

# ══════════════════════════════════════════════════
# 主界面
# ══════════════════════════════════════════════════
st.title("📚 OMNI PDF 电子书阅读工具")
st.markdown("本地大文件直接读取，无需上传 · DeepSeek AI 逐章总结 · 自动生成词汇表和练习题")

col_left, col_right = st.columns([1, 2])

# ── 左侧：文件选择 ──
with col_left:
    st.subheader("📁 选择 PDF 文件")

    # 方式1：本地路径输入
    st.markdown("**方式一：输入本地文件路径**")
    pdf_path_input = st.text_input(
        "PDF 文件路径",
        placeholder=r"例如：C:\Users\Administrator\Desktop\business_english.pdf",
        label_visibility="collapsed"
    )

    # 方式2：上传小文件
    st.markdown("**方式二：直接上传（<200MB）**")
    uploaded = st.file_uploader("上传 PDF", type=["pdf"],
                                label_visibility="collapsed")

    if uploaded:
        # 保存到临时文件
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.write(uploaded.getvalue())
        tmp.close()
        st.session_state.pdf_path = tmp.name
        st.session_state.pdf_name = uploaded.name
        st.success(f"✅ 已加载：{uploaded.name}")

    elif pdf_path_input and os.path.exists(pdf_path_input):
        st.session_state.pdf_path = pdf_path_input
        st.session_state.pdf_name = os.path.basename(pdf_path_input)
        st.success(f"✅ 已找到：{st.session_state.pdf_name}")

    elif pdf_path_input:
        st.error("❌ 文件路径不存在，请检查")

    # 设置
    if st.session_state.pdf_path:
        st.divider()
        st.subheader("⚙️ 处理设置")
        chunk_size = st.slider("每块页数", 3, 15, 6,
                               help="每次 AI 处理的页数，越少越精细但越慢")
        max_chunks = st.slider("最多处理章节数", 3, 20, 8,
                               help="限制处理范围，节省 API 费用")
        lang = st.selectbox("总结语言", ["中文", "English"], index=0)
        gen_questions = st.checkbox("同时生成练习题", value=True)

        st.divider()
        if st.button("🚀 开始 AI 分析", type="primary", use_container_width=True):
            st.session_state.summaries = {}
            st.session_state.questions = {}
            st.session_state.processing = True
            st.rerun()

    # 导出按钮
    if st.session_state.summaries:
        st.divider()
        st.subheader("💾 导出笔记")

        book_name = st.session_state.pdf_name.replace(".pdf", "")

        if st.button("📄 导出 Word 笔记", use_container_width=True):
            with st.spinner("生成 Word..."):
                path = export_to_word(book_name, st.session_state.summaries, {})
            if path:
                with open(path, "rb") as f:
                    st.download_button("⬇️ 下载 Word", f.read(),
                                       f"{book_name}_笔记.docx",
                                       use_container_width=True)

        if st.button("📊 导出 Excel 词汇表", use_container_width=True):
            with st.spinner("生成 Excel..."):
                path = export_to_excel(book_name, st.session_state.summaries)
            if path:
                with open(path, "rb") as f:
                    st.download_button("⬇️ 下载 Excel", f.read(),
                                       f"{book_name}_词汇表.xlsx",
                                       use_container_width=True)

# ── 右侧：内容展示 ──
with col_right:

    # 正在处理
    if st.session_state.processing and st.session_state.pdf_path:
        st.subheader("🤖 AI 正在分析中...")

        pages = load_pdf(st.session_state.pdf_path)
        if not pages:
            st.error("PDF 读取失败")
            st.session_state.processing = False
        else:
            st.session_state.total_pages = len(pages)
            st.info(f"📄 共 {len(pages)} 页，开始逐块分析...")

            chunks = pages_to_chunks(pages, chunk_size)[:max_chunks]
            book_name = st.session_state.pdf_name

            progress_bar = st.progress(0)
            status = st.empty()

            for i, chunk in enumerate(chunks):
                chunk_key = f"第 {chunk['start_page']}-{chunk['end_page']} 页"
                status.markdown(f'<div class="progress-info">🔍 正在分析 {chunk_key}...</div>',
                               unsafe_allow_html=True)

                lang_code = "zh" if lang == "中文" else "en"
                summary = summarize_chunk(chunk["text"], book_name, lang_code)
                st.session_state.summaries[chunk_key] = summary

                if gen_questions:
                    qs = generate_questions_for_chunk(chunk["text"], book_name)
                    st.session_state.questions[chunk_key] = qs

                progress_bar.progress((i + 1) / len(chunks))
                time.sleep(0.3)

            status.empty()
            progress_bar.empty()
            st.session_state.processing = False
            st.success(f"✅ 分析完成！处理了 {len(chunks)} 个章节")
            st.rerun()

    # 显示结果
    elif st.session_state.summaries:
        # Tab 选择
        tab_summary, tab_vocab, tab_quiz = st.tabs([
            "📝 知识点总结", "🔑 词汇表", "❓ 练习题"
        ])

        with tab_summary:
            st.subheader(f"📖 {st.session_state.pdf_name}")
            st.caption(f"共 {st.session_state.total_pages} 页 · "
                      f"已分析 {len(st.session_state.summaries)} 个章节")

            for chunk_key, summary in st.session_state.summaries.items():
                with st.expander(f"📄 {chunk_key}", expanded=False):
                    if summary.get("key_points"):
                        st.markdown("**🎯 核心知识点**")
                        for pt in summary["key_points"]:
                            st.markdown(f"• {pt}")

                    if summary.get("useful_phrases"):
                        st.markdown("**💬 实用短语**")
                        for ph in summary["useful_phrases"]:
                            st.markdown(f"› `{ph}`")

        with tab_vocab:
            st.subheader("🔑 全书词汇汇总")
            all_vocab = []
            for chunk_key, summary in st.session_state.summaries.items():
                for v in summary.get("vocabulary", []):
                    v["chapter"] = chunk_key
                    all_vocab.append(v)

            st.caption(f"共提取 {len(all_vocab)} 个词汇")

            # 搜索
            search = st.text_input("🔍 搜索词汇", placeholder="输入英文或中文搜索...")
            filtered = [v for v in all_vocab if
                       search.lower() in v.get("word","").lower() or
                       search in v.get("cn","")] if search else all_vocab

            for v in filtered:
                st.markdown(
                    f"<span class='vocab-pill'>{v.get('word','')}</span> "
                    f"— {v.get('cn','')}",
                    unsafe_allow_html=True
                )
                if v.get("example"):
                    st.caption(f"  例：{v.get('example','')}")

        with tab_quiz:
            st.subheader("❓ 练习题")
            if not st.session_state.questions:
                st.info("生成时请勾选「同时生成练习题」")
            else:
                for chunk_key, qs in st.session_state.questions.items():
                    if not qs:
                        continue
                    st.markdown(f"**📄 {chunk_key}**")
                    for qi, q in enumerate(qs, 1):
                        qtype = q.get("type","")
                        with st.expander(f"Q{qi}. {q.get('question','')[:60]}...",
                                        expanded=False):
                            st.write(q.get("question",""))
                            opts = q.get("options",[])
                            if opts:
                                for o in opts:
                                    st.write(o)
                            with st.expander("查看答案"):
                                st.success(f"答案：{q.get('answer','')}")
                                st.caption(q.get("explanation",""))

    elif not st.session_state.pdf_path:
        st.info("👈 在左侧输入 PDF 文件路径或上传文件，开始 AI 分析。")
        st.markdown("""
**支持两种方式加载大型 PDF：**

**方式一（推荐，无文件大小限制）：**
复制 PDF 在电脑上的完整路径，粘贴到左侧输入框。

Windows 路径示例：
```
C:\\Users\\Administrator\\Desktop\\business_english.pdf
```

**方式二：**
直接上传 200MB 以内的 PDF 文件。

---

**分析内容包括：**
- 📝 每章核心知识点总结
- 🔑 重要词汇（含中文释义和例句）  
- 💬 实用短语和表达
- ❓ 配套练习题
- 💾 一键导出 Word 笔记 / Excel 词汇表
        """)
